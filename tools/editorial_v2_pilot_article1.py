from __future__ import annotations

from pathlib import Path

from docx import Document


BASE = Path(r"E:\Carte\BB\++++carti scrise de bebe\CELE 63 de calitati ale liderului")
SOURCE = BASE / "pentru tiparire - actualizat cu articole web.docx"
OUT = BASE / "Final Corectat V2 - pilot articol 1.docx"
REPORT = BASE / "raport-corectura-editoriala-v2-pilot.txt"


OLD_DIACRITICS = str.maketrans({"ş": "ș", "Ş": "Ș", "ţ": "ț", "Ţ": "Ț"})


ARTICLE_1 = {
    168: "Viziunea bazată pe curiozitate și explorare",
    169: "Dezvoltă-ți potențialul de lider construind o viziune pe măsura propriilor tale idei și experiențe.",
    171: (
        "Fiindcă am fost mereu atras de biografiile oamenilor celebri, am citit de curând "
        "„Finita Comedia”, în care autorul, Romain Rolland, prezintă numeroase momente din viața "
        "marelui compozitor german Ludwig van Beethoven. Așa am aflat despre unul dintre marile "
        "secrete ale lui Beethoven, care stăteau la baza realizării compozițiilor sale."
    ),
    172: (
        "Tânărul valet, Michael Krenn, relata cum își întrebuința marele compozitor timpul: "
        "la 05:30 se trezea și se apuca de lucru, cu mare gălăgie, cântând, strigând, dând "
        "din picioare și din mâini ca să marcheze măsura. La 07:30 își lua micul dejun, apoi "
        "îndată o pornea pe câmp; părea un nebun care vorbea tare, vocifera, își mișca brațele "
        "ca o moară, fără să se sinchisească de țăranii care-i strigau că le calcă ogoarele și "
        "le sperie turmele. La 12:30 se întorcea pentru prânz, se urca la el în cameră până la "
        "15:00, apoi o lua din nou razna pe câmp. La apusul soarelui se întorcea. La 19:30 cina "
        "și se închidea după aceea la el în cameră; compunea până la 22:00 și se băga cuminte în pat."
    ),
    173: "Experiența este cea care dă avânt viziunii tale",
    174: (
        "Pentru Beethoven, acele frumoase plimbări în aer liber erau rodnice pentru spirit. "
        "Pășea în acea minune numită natură și o trăia cu toate cele cinci simțuri. Nu putea "
        "să lase muza să doarmă. Nădăjduia să aducă lumii mereu opere noi, cât mai valoroase. "
        "Nu este de mirare că viața lui era mereu activă."
    ),
    175: (
        "Fără umbră de îndoială, lui Beethoven i se părea mai ușor să creeze o nouă operă "
        "petrecându-și timpul în acel ambient unic, magic, creat chiar de el însuși - plin de "
        "zgomote, sunete, mirosuri, mișcări și imagini. Deseori făcea o impresie proastă, fiindcă "
        "nimeni nu-l înțelegea. Dar numai el, un adevărat maestru al compoziției, găsea în toate "
        "acele obiceiuri o sursă incredibilă de inspirație. Trebuia să perceapă toate detaliile și "
        "dinamica semnalelor acustice și vizuale, pentru a-și putea crea în minte o imagine despre "
        "cum avea să arate viitoarea partitură. Asocia fiecare sunet cu o anumită trăire și cu un "
        "anumit sentiment, pentru a crea o armonie."
    ),
    176: "Avea mare dreptate cine afirma că viziunea este acea „imagine” puternică ce declanșează emoții și entuziasm în sufletul tău.",
    177: "Dă mai multă viață experiențelor, dacă vrei să creezi o viziune care să te uluiască",
    178: (
        "Experiența pe care o poți fructifica în mod concret prin diverse activități pe care le "
        "desfășori, acasă sau la locul de muncă, te poate ajuta să-ți creezi o viziune. Iar ideile "
        "pe care le descoperi te pot ajuta să abordezi lucrurile dintr-o altă perspectivă. Totul "
        "în activitățile tale trebuie să se transforme într-un joc interesant de căutare a „muzei” "
        "care să te ajute să compui cu pasiune propria ta „partitură”."
    ),
    179: (
        "Identificarea acelor lucruri, acțiuni, situații sau particularități care te fac să te simți "
        "bine în forul tău interior devine un mijloc de expresie a propriei tale „compoziții”. Pune-ți "
        "întotdeauna întrebarea: ce a avut deosebit acea experiență, încât să te impresioneze? Dacă "
        "viziunea nu te impresionează, atunci nu este suficient de reală pentru a putea fi pusă în "
        "aplicare. Ea trebuie să fie legată de sufletul tău, să te incite, să te atragă tot mai mult, "
        "să-ți trezească interesul și să-ți dea ocazia să-ți pui ideile în practică."
    ),
    180: "Tu ai o imagine a viziunii pe care încerci să o creezi? Cum creezi această imagine?",
    181: (
        "Așa cum poezia, ca să fie pe deplin simțită și înțeleasă, trebuie recitată cu o intonație "
        "adecvată, tot așa și viziunea ta, ca să poată fi transpusă în realitate, trebuie să fie "
        "accesibilă și clară, să vină din sufletul tău, să fie flexibilă și mereu adaptabilă și să "
        "exprime un anumit ideal."
    ),
    182: (
        "Viziunea ta trebuie să răspundă la următoarea întrebare: „Ce anume dorești să realizezi?” "
        "Nu uita că, dacă Beethoven n-ar fi avut un scop precis, plimbările sale n-ar mai fi avut niciun sens."
    ),
    183: (
        "Indiferent că ești CEO sau muzician, trebuie să-ți creezi propriile momente de formare a "
        "viziunii, plecând de la evenimentele petrecute în viața ta. Propriile tale experiențe zilnice, "
        "trăite în compania altor oameni, îți pot furniza o formulă magică prin care să schimbi lumea."
    ),
    184: (
        "O viziune care să te încurajeze, să te sprijine în demersurile tale și să-ți conducă deciziile "
        "în direcția corectă este greu de imaginat, dar și mai greu de acceptat. Nu toate gândurile care "
        "îți vin în minte îți pot fi de folos, la fel cum nu toate semințele sunt roditoare."
    ),
    185: (
        "Concluzie: Reușita ta, ca lider, depinde foarte mult de felul în care vezi lucrurile. Nu vei "
        "reuși să-ți atingi obiectivele dacă nu ai o viziune de ansamblu asupra propriei tale „opere”, "
        "plină de speranță și de idei, și dacă nu realizezi un plan viabil care să continue preocupările "
        "tale dintr-o perspectivă mai clară. Dacă reușești să-ți creezi o viziune favorabilă, plecând de "
        "la experiențele zilnice pe care le trăiești, îți vei putea spune la sfârșitul fiecărei zile: "
        "„Timpul meu a fost bine petrecut azi.”"
    ),
    186: "Dezvoltă-ți potențialul de lider construind o viziune pe măsura propriilor tale idei și experiențe.",
}


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main() -> None:
    doc = Document(SOURCE)

    changed = []
    for idx, text in ARTICLE_1.items():
        old = doc.paragraphs[idx].text
        set_paragraph_text(doc.paragraphs[idx], text)
        changed.append((idx, old.translate(OLD_DIACRITICS), text))

    doc.save(OUT)

    lines = [
        "Corectura editoriala V2 - pilot articol 1",
        f"Sursa: {SOURCE}",
        f"Rezultat: {OUT}",
        "",
        "Schimbari aplicate:",
    ]
    for idx, old, new in changed:
        lines.append(f"\nParagraf DOCX {idx}")
        lines.append(f"INAINTE: {old}")
        lines.append(f"DUPA:    {new}")
    REPORT.write_text("\n".join(lines), encoding="utf-8")

    print(OUT)
    print(REPORT)
    print(f"paragraphs_changed={len(changed)}")


if __name__ == "__main__":
    main()
