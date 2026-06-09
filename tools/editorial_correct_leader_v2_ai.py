from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from openai_key import get_openai_api_key


BASE_DIR = Path(r"E:\Carte\BB\++++carti scrise de bebe\CELE 63 de calitati ale liderului")
TOOLS_DIR = Path(r"D:\Teste cursor\HTML Editor\Word Editor\tools")
SOURCE_DOCX = BASE_DIR / "pentru tiparire.docx"
OUT_DOCX = BASE_DIR / "Final Corectat V2 GPT-5.4.docx"
REPORT_TXT = BASE_DIR / "raport-corectura-editoriala-v2-ai.txt"
CACHE_JSONL = TOOLS_DIR / "editorial_v2_ai_cache.jsonl"

API_URL = "https://api.openai.com/v1/responses"
PROMPT_VERSION = "v2.2-punctuation-final-review"

# === CHEIA OPENAI API ===
# Pune cheia o singura data in:
# D:\Teste cursor\HTML Editor\Word Editor\tools\openai_api_key.txt
# Toate scripturile Python o pot citi prin tools\openai_key.py.

MODELS_BY_PREFERENCE = [
    "gpt-5.4",
    "gpt-5.5",
    "gpt-5.4-mini",
    "gpt-4.1",
    "gpt-4o",
]


SYSTEM_PROMPT = """
Esti corector editorial profesionist pentru limba romana.

Vei corecta fragmente dintr-o carte despre leadership. Lucreaza cu fidelitate maxima fata de stilul autorului:
- corecteaza gramatica, ortografia, diacriticele, punctuatia, acordurile, topica greoaie, logica frazei si exprimarile nenaturale;
- pastreaza intentia, tonul si vocabularul autorului cat mai mult posibil;
- nu rescrie creativ si nu moderniza textul daca fraza are deja sens;
- adauga cuvinte doar cand lipsesc gramatical sau logic;
- schimba un cuvant numai cand cel original suna gresit, fortat, confuz sau nenatural;
- nu rezuma, nu extinde idei, nu inventa exemple si nu elimina informatie;
- pastreaza nume, titluri, citate, cifre, ISBN, ani, marcaje de articol si sensul exact;
- pastreaza fiecare paragraf separat: primesti N paragrafe si returnezi exact N paragrafe, cu aceleasi id-uri.

Revizuire finala obligatorie inainte de raspuns:
- dupa ce corectezi fiecare paragraf, reciteste-l inca o data strict pentru virgule, citate, acorduri si fluenta;
- verifica punctuatia dinaintea citatelor sau a vorbirii directe: de exemplu, "iti vei putea spune, „Timpul meu...”";
- verifica virgulele care clarifica subordonate, incidente, enumerari si treceri intre idei;
- verifica daca termenul ales este cel mai natural in context, dar fara rescrieri inutile;
- nu returna explicatii, ci doar JSON-ul final cu paragrafele deja verificate.

Calibrare dupa corectura dorita de autor:
Original: "Asa cum doctorii trebuie sa descalceasca firele incurcate ale bolii, sa stie ce intrebari sa puna pacientilor, sa recunoasca indiciile fizice subtile si sa identifice testele care ar putea duce, in cele din urma, la diagnosticul perfect - tot asa si tu trebuie sa descalcești modul in care poti obtine rezultatele cele mai bune..."
Corect: "Asa cum doctorii trebuie sa descalceasca firele incurcate ale bolii, sa stie ce intrebari sa le puna pacientilor, sa recunoasca indiciile fizice subtile si sa identifice testele care ar putea duce, in cele din urma, la diagnosticul corect - tot asa si tu trebuie sa descalcesti modul in care poti obtine cele mai bune rezultate..."

Original: "Daca vrei sa fii o persoana care sa-i conduca pe altii si sa indrume o organizatie..."
Corect: "Daca vrei sa fii o persoana capabila sa-i conduca pe altii si sa indrume o organizatie..."

Returneaza numai JSON valid, fara markdown:
{"paragraphs":[{"id":123,"text":"text corectat"}]}
""".strip()


@dataclass
class ParagraphItem:
    idx: int
    text: str


@dataclass
class ApiResult:
    texts: dict[int, str]
    usage: dict[str, int]


GPT54_INPUT_USD_PER_1M = 2.50
GPT54_OUTPUT_USD_PER_1M = 15.00


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def normalize_old_romanian_diacritics(text: str) -> str:
    return text.translate(str.maketrans({"ş": "ș", "Ş": "Ș", "ţ": "ț", "Ţ": "Ț"}))


def compact_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def ensure_source_docx() -> None:
    if SOURCE_DOCX.is_file():
        return
    raise FileNotFoundError(f"Lipseste sursa {SOURCE_DOCX}")


def is_editable_paragraph(text: str) -> bool:
    t = compact_text(text)
    if not t:
        return False
    if re.fullmatch(r"-\s*\d+\s*-", t):
        return False
    if re.fullmatch(r"\d+\.?", t):
        return False
    if len(t) <= 2:
        return False
    return True


def collect_items(doc: Document) -> list[ParagraphItem]:
    items: list[ParagraphItem] = []
    for idx, par in enumerate(doc.paragraphs):
        text = normalize_old_romanian_diacritics(par.text)
        if is_editable_paragraph(text):
            items.append(ParagraphItem(idx=idx, text=text))
    return items


def find_chapter_bounds(doc: Document, chapter: int) -> tuple[int, int]:
    starts: list[tuple[int, int]] = []
    for idx, par in enumerate(doc.paragraphs):
        text = compact_text(par.text)
        if re.fullmatch(rf"-\s*{chapter}\s*-", text):
            starts.append((chapter, idx))
        elif re.fullmatch(r"-\s*\d+\s*-", text):
            m = re.fullmatch(r"-\s*(\d+)\s*-", text)
            if m:
                starts.append((int(m.group(1)), idx))
    starts.sort(key=lambda item: item[1])
    for pos, (num, idx) in enumerate(starts):
        if num == chapter:
            end = starts[pos + 1][1] if pos + 1 < len(starts) else len(doc.paragraphs)
            return idx + 1, end
    raise ValueError(f"Nu gasesc capitolul {chapter} in document")


def filter_items_by_range(items: list[ParagraphItem], start: int | None, end: int | None) -> list[ParagraphItem]:
    if start is None or end is None:
        return items
    return [item for item in items if start <= item.idx < end]


def make_batches(items: list[ParagraphItem], max_chars: int = 5200, max_items: int = 18) -> list[list[ParagraphItem]]:
    batches: list[list[ParagraphItem]] = []
    cur: list[ParagraphItem] = []
    chars = 0
    for item in items:
        need = len(item.text) + 80
        if cur and (chars + need > max_chars or len(cur) >= max_items):
            batches.append(cur)
            cur = []
            chars = 0
        cur.append(item)
        chars += need
    if cur:
        batches.append(cur)
    return batches


def read_cache() -> dict[str, dict[int, str]]:
    cache: dict[str, dict[int, str]] = {}
    if not CACHE_JSONL.is_file():
        return cache
    with CACHE_JSONL.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                rec = json.loads(line)
                cache[rec["batch_hash"]] = {int(k): v for k, v in rec["texts"].items()}
            except Exception:
                continue
    return cache


def append_cache(batch_hash: str, texts: dict[int, str], model: str) -> None:
    rec = {
        "batch_hash": batch_hash,
        "model": model,
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "texts": {str(k): v for k, v in texts.items()},
    }
    with CACHE_JSONL.open("a", encoding="utf-8") as f:
        f.write(json.dumps(rec, ensure_ascii=False) + "\n")


def batch_hash(batch: list[ParagraphItem], model: str, temperature: float | None) -> str:
    payload = {
        "prompt_version": PROMPT_VERSION,
        "model": model,
        "temperature": temperature,
        "paragraphs": [{"id": it.idx, "text": it.text} for it in batch],
    }
    return sha256_text(json.dumps(payload, ensure_ascii=False, sort_keys=True))


def extract_response_text(data: dict[str, Any]) -> str:
    if "output_text" in data:
        return str(data["output_text"])
    parts: list[str] = []
    for item in data.get("output", []):
        for content in item.get("content", []):
            if content.get("type") in {"output_text", "text"} and "text" in content:
                parts.append(str(content["text"]))
    return "\n".join(parts)


def extract_usage(data: dict[str, Any]) -> dict[str, int]:
    usage = data.get("usage") or {}
    input_tokens = int(usage.get("input_tokens") or usage.get("prompt_tokens") or 0)
    output_tokens = int(usage.get("output_tokens") or usage.get("completion_tokens") or 0)
    total_tokens = int(usage.get("total_tokens") or (input_tokens + output_tokens))
    return {
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "total_tokens": total_tokens,
    }


def estimate_gpt54_cost(usage: dict[str, int]) -> float:
    input_tokens = usage.get("input_tokens", 0)
    output_tokens = usage.get("output_tokens", 0)
    return (input_tokens / 1_000_000 * GPT54_INPUT_USD_PER_1M) + (output_tokens / 1_000_000 * GPT54_OUTPUT_USD_PER_1M)


def call_openai(batch: list[ParagraphItem], model: str, temperature: float | None, timeout: int = 180) -> ApiResult:
    key = get_openai_api_key()

    payload = {
        "model": model,
        "input": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": json.dumps(
                    {
                        "task": "Corecteaza editorial aceste paragrafe si returneaza exact JSON-ul cerut.",
                        "paragraphs": [{"id": it.idx, "text": it.text} for it in batch],
                    },
                    ensure_ascii=False,
                ),
            },
        ],
        "text": {"format": {"type": "json_object"}},
    }
    if temperature is not None:
        payload["temperature"] = temperature
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        API_URL,
        data=body,
        headers={
            "Authorization": f"Bearer {key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        data = json.loads(resp.read().decode("utf-8"))

    text = extract_response_text(data).strip()
    usage = extract_usage(data)
    parsed = json.loads(text)
    returned = parsed.get("paragraphs")
    if not isinstance(returned, list):
        raise ValueError("Raspuns fara lista paragraphs")
    out: dict[int, str] = {}
    expected = {it.idx for it in batch}
    for row in returned:
        idx = int(row["id"])
        if idx not in expected:
            raise ValueError(f"ID neasteptat in raspuns: {idx}")
        out[idx] = str(row.get("text", ""))
    if set(out) != expected:
        missing = sorted(expected - set(out))
        raise ValueError(f"Lipsesc paragrafe din raspuns: {missing[:10]}")
    return ApiResult(texts=out, usage=usage)


def corrected_text_is_suspicious(old: str, new: str) -> bool:
    old_c = compact_text(old)
    new_c = compact_text(new)
    if not new_c and old_c:
        return True
    if "```" in new_c or new_c.startswith("{") or new_c.startswith("["):
        return True
    if len(old_c) > 80:
        if len(new_c) < len(old_c) * 0.45:
            return True
        if len(new_c) > len(old_c) * 1.75:
            return True
    return False


def choose_model(explicit_model: str | None) -> str:
    if explicit_model:
        return explicit_model
    env_model = os.environ.get("OPENAI_MODEL", "").strip()
    if env_model:
        return env_model
    return MODELS_BY_PREFERENCE[0]


def set_paragraph_text(paragraph, text: str) -> None:
    # Pastreaza stilul paragrafului si proprietatile primului run; evita refacerea structurii DOCX.
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def write_report(
    source: Path,
    out: Path,
    model: str,
    temperature: float | None,
    items: list[ParagraphItem],
    corrected: dict[int, str],
    changed: list[tuple[int, str, str]],
    limit: int | None,
    usage_total: dict[str, int],
    estimated_cost: float,
    cached_batches: int,
    api_batches: int,
) -> None:
    lines = [
        "Corectura editoriala V2 AI",
        f"Sursa: {source}",
        f"Rezultat: {out}",
        f"Model: {model}",
        f"Temperature: {temperature if temperature is not None else 'default'}",
        f"Paragrafe eligibile: {len(items)}",
        f"Paragrafe corectate in aceasta rulare: {len(corrected)}",
        f"Paragrafe schimbate: {len(changed)}",
        f"Batch-uri din cache: {cached_batches}",
        f"Batch-uri API facturabile in aceasta rulare: {api_batches}",
        f"Input tokens API: {usage_total.get('input_tokens', 0)}",
        f"Output tokens API: {usage_total.get('output_tokens', 0)}",
        f"Total tokens API: {usage_total.get('total_tokens', 0)}",
        f"Cost estimat GPT-5.4: ${estimated_cost:.4f}",
    ]
    if limit:
        lines.append(f"LIMITA TEST: {limit} paragrafe")
    lines.extend(["", "Exemple de schimbari:"])
    for idx, old, new in changed[:120]:
        lines.append("")
        lines.append(f"Paragraf DOCX {idx}")
        lines.append(f"INAINTE: {old}")
        lines.append(f"DUPA:    {new}")
    REPORT_TXT.write_text("\n".join(lines), encoding="utf-8-sig")


def run(args: argparse.Namespace) -> None:
    ensure_source_docx()
    doc = Document(SOURCE_DOCX)
    items = collect_items(doc)
    range_start = range_end = None
    if args.chapter:
        range_start, range_end = find_chapter_bounds(doc, args.chapter)
        items = filter_items_by_range(items, range_start, range_end)
    if args.limit:
        items = items[: args.limit]
    batches = make_batches(items, max_chars=args.max_chars, max_items=args.max_items)
    cache = read_cache()
    model = choose_model(args.model)
    out_docx = Path(args.out) if args.out else OUT_DOCX
    if args.chapter and not args.out:
        out_docx = BASE_DIR / f"Final Corectat V2 GPT-5.4 - test capitol {args.chapter}.docx"

    corrected: dict[int, str] = {}
    usage_total = {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}
    cached_batches = 0
    api_batches = 0
    for n, batch in enumerate(batches, start=1):
        bh = batch_hash(batch, model, args.temperature)
        if not args.no_cache and bh in cache:
            corrected.update(cache[bh])
            cached_batches += 1
            print(f"[{n}/{len(batches)}] cache {len(batch)} paragrafe")
            continue

        print(f"[{n}/{len(batches)}] OpenAI {model}: {len(batch)} paragrafe")
        last_error: Exception | None = None
        for attempt in range(1, args.retries + 1):
            try:
                result = call_openai(batch, model=model, temperature=args.temperature, timeout=args.timeout)
                out = result.texts
                for item in batch:
                    new = normalize_old_romanian_diacritics(out[item.idx])
                    if corrected_text_is_suspicious(item.text, new):
                        raise ValueError(f"Corectura suspecta la paragraful {item.idx}")
                    out[item.idx] = new
                corrected.update(out)
                api_batches += 1
                for key in usage_total:
                    usage_total[key] += result.usage.get(key, 0)
                if not args.no_cache:
                    append_cache(bh, out, model)
                break
            except urllib.error.HTTPError as e:
                detail = e.read().decode("utf-8", "replace")
                last_error = RuntimeError(f"HTTP {e.code}: {detail[:500]}")
                if e.code == 429 and "insufficient_quota" in detail:
                    raise RuntimeError(
                        "API key-ul este valid, dar proiectul OpenAI nu are quota/credit disponibil "
                        "(insufficient_quota). Activeaza billing sau mareste limita proiectului, apoi ruleaza din nou."
                    )
                if e.code in {401, 403, 404}:
                    raise last_error
            except Exception as e:
                last_error = e
            wait = min(8 * attempt, 45)
            print(f"  incercare {attempt} esuata: {last_error}; astept {wait}s")
            time.sleep(wait)
        else:
            raise RuntimeError(f"Nu pot corecta batch-ul {n}: {last_error}")

    changed: list[tuple[int, str, str]] = []
    for item in items:
        new = corrected.get(item.idx)
        if new is None:
            continue
        old = normalize_old_romanian_diacritics(doc.paragraphs[item.idx].text)
        if compact_text(old) != compact_text(new):
            changed.append((item.idx, old, new))
        if not args.dry_run:
            set_paragraph_text(doc.paragraphs[item.idx], new)

    if not args.dry_run:
        doc.save(out_docx)
    estimated_cost = estimate_gpt54_cost(usage_total)
    write_report(SOURCE_DOCX, out_docx, model, args.temperature, items, corrected, changed, args.limit, usage_total, estimated_cost, cached_batches, api_batches)
    print(f"source={SOURCE_DOCX}")
    print(f"out={out_docx}")
    print(f"report={REPORT_TXT}")
    print(f"eligible={len(items)} corrected={len(corrected)} changed={len(changed)}")
    print(f"api_batches={api_batches} cached_batches={cached_batches}")
    print(f"input_tokens={usage_total.get('input_tokens', 0)} output_tokens={usage_total.get('output_tokens', 0)} total_tokens={usage_total.get('total_tokens', 0)}")
    print(f"estimated_cost_gpt54_usd={estimated_cost:.4f}")
    if args.chapter:
        print(f"chapter={args.chapter} range={range_start}:{range_end}")
    if args.dry_run:
        print("DRY_RUN: documentul DOCX nu a fost scris")


def main() -> None:
    parser = argparse.ArgumentParser(description="Corectura editoriala AI V2 pentru cartea de leadership.")
    parser.add_argument("--model", default=None, help="Model OpenAI; implicit OPENAI_MODEL sau gpt-5.4")
    parser.add_argument("--temperature", type=float, default=0.1, help="Temperatura modelului; foloseste --temperature -1 pentru default fara parametru.")
    parser.add_argument("--chapter", type=int, default=None, help="Corecteaza doar capitolul indicat, dupa marcajele -1-, -2- etc.")
    parser.add_argument("--out", default=None, help="Calea DOCX de iesire.")
    parser.add_argument("--limit", type=int, default=None, help="Proceseaza doar primele N paragrafe eligibile.")
    parser.add_argument("--max-chars", type=int, default=5200)
    parser.add_argument("--max-items", type=int, default=18)
    parser.add_argument("--timeout", type=int, default=180)
    parser.add_argument("--retries", type=int, default=3)
    parser.add_argument("--no-cache", action="store_true")
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    if args.temperature is not None and args.temperature < 0:
        args.temperature = None
    run(args)


if __name__ == "__main__":
    main()
