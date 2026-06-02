from __future__ import annotations

import copy
import re
import unicodedata
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import html


DOCX_PATH = Path(
    r"e:\Carte\BB\++++carti scrise de bebe\CELE 63 de calitati ale liderului\pentru tiparire.docx"
)
HTML_ROOT = Path(r"e:\Carte\BB\17 - Site Leadership\Principal\ro")
OUT_DOCX = DOCX_PATH.with_name("pentru tiparire - actualizat cu articole web.docx")
REPORT_PATH = DOCX_PATH.with_name("raport-inlocuiri-web.txt")

HTML_FILES = [
    "calitatile-unui-lider-inspiratia.html",
    "calitatile-unui-lider-responsabilitatea.html",
    "calitatile-unui-lider-credinta.html",
    "calitatile-unui-lider-dorinta-de-autodepasire.html",
    "calitatile-unui-lider-increderea.html",
    "calitatile-unui-lider-perseverenta.html",
    "calitatile-unui-lider-spontaneitatea.html",
    "calitatile-unui-lider-vointa-ferma-de-a-invinge.html",
    "calitatile-unui-lider-rezonanta.html",
    "calitatile-unui-lider-sensibilitatea-sufleteasca.html",
    "calitatile-unui-lider-receptivitatea.html",
    "calitatile-unui-lider-puterea-de-patrundere-psihologica.html",
    "calitatile-unui-lider-puterea-de-persuasiune.html",
    "calitatile-unui-lider-maretia-sufleteasca.html",
    "calitatile-unui-lider-puterea-de-patrundere-a-eu-lui-individual.html",
    "calitatile-unui-lider-luciditatea-si-profunzimea-judecatii.html",
    "calitatile-unui-lider-maretia-spirituala.html",
    "calitatile-unui-lider-integritatea-launtrica.html",
    "calitatile-unui-lider-intuitia-si-viziunea-patrunzatoare.html",
    "calitatile-unui-lider-forta-emotionala.html",
    "calitatile-unui-lider-devotamentul-absolut.html",
    "calitatile-unui-lider-carisma.html",
    "calitatile-unui-lider-creativitatea.html",
]


@dataclass
class HtmlPara:
    kind: str
    text: str


@dataclass
class HtmlArticle:
    file: Path
    title: str
    subtitle: str
    paras: list[HtmlPara]


@dataclass
class DocArticle:
    number: int
    marker_idx: int
    end_idx: int
    title: str
    subtitle: str


def romanian_key(text: str) -> str:
    text = text.strip().lower()
    text = text.replace("ȋ", "î").replace("Ȋ", "î")
    text = text.replace("ş", "ș").replace("ţ", "ț")
    text = text.replace("ş", "ș").replace("ţ", "ț")
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = text.replace("ş", "s").replace("ț", "t").replace("ţ", "t")
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def parse_html_article(path: Path) -> HtmlArticle:
    tree = html.fromstring(path.read_bytes())
    body = tree.xpath('//div[@itemprop="articleBody"]')
    if not body:
        raise RuntimeError(f"Nu gasesc articleBody in {path}")
    body = body[0]
    title = clean_text(body.xpath("string(.//h1[1])"))
    paras: list[HtmlPara] = []
    for p in body.xpath('.//p[contains(@class, "text_obisnuit")]'):
        text = clean_text("".join(p.itertext()))
        if not text:
            continue
        if text.lower().startswith("ultimele articole"):
            break
        cls = p.get("class") or ""
        kind = "lead" if "text_obisnuit2" in cls else "body"
        paras.append(HtmlPara(kind=kind, text=text))
    subtitle = paras[0].text if paras else ""
    return HtmlArticle(file=path, title=title, subtitle=subtitle, paras=paras)


def find_doc_articles(doc: Document) -> list[DocArticle]:
    starts: list[tuple[int, int]] = []
    for idx, par in enumerate(doc.paragraphs):
        text = clean_text(par.text)
        m = re.match(r"^-\s*(\d+)\s*-$", text)
        if m:
            starts.append((int(m.group(1)), idx))

    articles: list[DocArticle] = []
    for pos, (number, idx) in enumerate(starts):
        end = starts[pos + 1][1] if pos + 1 < len(starts) else len(doc.paragraphs)
        nonempty = []
        for j in range(idx + 1, min(end, idx + 10)):
            text = clean_text(doc.paragraphs[j].text)
            if text:
                nonempty.append((j, text))
            if len(nonempty) >= 2:
                break
        title = nonempty[0][1] if nonempty else ""
        subtitle = nonempty[1][1] if len(nonempty) > 1 else ""
        articles.append(DocArticle(number=number, marker_idx=idx, end_idx=end, title=title, subtitle=subtitle))
    return articles


def para_has_text(par: Paragraph) -> bool:
    return bool(clean_text(par.text))


def delete_paragraph(par: Paragraph) -> None:
    el = par._element
    el.getparent().remove(el)


def text_template(par: Paragraph):
    el = par._element
    first_run = el.find(qn("w:r"))
    r_pr = None
    if first_run is not None:
        r_pr_el = first_run.find(qn("w:rPr"))
        if r_pr_el is not None:
            r_pr = copy.deepcopy(r_pr_el)
    return copy.deepcopy(el), r_pr


def set_p_text(p_el, text: str, r_pr=None):
    for child in list(p_el):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            p_el.remove(child)
    r = OxmlElement("w:r")
    if r_pr is not None:
        r.append(copy.deepcopy(r_pr))
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p_el.append(r)


def insert_after(prev_el, template_el, text: str, r_pr=None):
    new_el = copy.deepcopy(template_el)
    set_p_text(new_el, text, r_pr)
    prev_el.addnext(new_el)
    return new_el


def article_templates(doc: Document, article: DocArticle):
    title_par = None
    subtitle_par = None
    body_par = None
    bold_body_par = None
    blank_par = None
    nonempty = []
    for idx in range(article.marker_idx + 1, article.end_idx):
        par = doc.paragraphs[idx]
        if para_has_text(par):
            nonempty.append(par)
        elif blank_par is None:
            blank_par = par
    if nonempty:
        title_par = nonempty[0]
    if len(nonempty) > 1:
        subtitle_par = nonempty[1]
    for par in nonempty[2:]:
        if body_par is None:
            body_par = par
        if bold_body_par is None and any(run.bold for run in par.runs):
            bold_body_par = par
    body_par = body_par or subtitle_par or title_par or doc.paragraphs[article.marker_idx]
    bold_body_par = bold_body_par or body_par
    blank_par = blank_par or body_par
    return {
        "title": text_template(title_par or body_par),
        "subtitle": text_template(subtitle_par or body_par),
        "body": text_template(body_par),
        "bold": text_template(bold_body_par),
        "blank": text_template(blank_par),
    }


def best_doc_match(html_article: HtmlArticle, doc_articles: list[DocArticle], used_numbers: set[int]):
    hkey = romanian_key(html_article.title)
    by_key = {romanian_key(a.title): a for a in doc_articles}
    if hkey in by_key and by_key[hkey].number not in used_numbers:
        return by_key[hkey], 1.0, "exact"

    candidates = []
    for article in doc_articles:
        if article.number in used_numbers:
            continue
        score = SequenceMatcher(None, hkey, romanian_key(article.title)).ratio()
        candidates.append((score, article))
    score, article = max(candidates, key=lambda x: x[0])
    if score >= 0.62:
        return article, score, "fuzzy"
    return None, score, "unmatched"


def replace_article(doc: Document, article: DocArticle, html_article: HtmlArticle):
    templates = article_templates(doc, article)
    marker_el = doc.paragraphs[article.marker_idx]._element
    to_delete = [doc.paragraphs[i] for i in range(article.marker_idx + 1, article.end_idx)]
    for par in to_delete:
        delete_paragraph(par)

    prev = marker_el
    title_el, title_rpr = templates["title"]
    subtitle_el, subtitle_rpr = templates["subtitle"]
    body_el, body_rpr = templates["body"]
    bold_el, bold_rpr = templates["bold"]
    blank_el, blank_rpr = templates["blank"]

    prev = insert_after(prev, title_el, html_article.title, title_rpr)
    if html_article.paras:
        prev = insert_after(prev, subtitle_el, html_article.paras[0].text, subtitle_rpr)
        prev = insert_after(prev, blank_el, "", blank_rpr)
        for hp in html_article.paras[1:]:
            if hp.kind == "lead":
                prev = insert_after(prev, bold_el, hp.text, bold_rpr)
            else:
                prev = insert_after(prev, body_el, hp.text, body_rpr)


def append_unmatched(doc: Document, unmatched: list[HtmlArticle]):
    if not unmatched:
        return
    doc.add_page_break()
    h = doc.add_paragraph("Articole de pe site fără corespondent sigur în carte")
    h.style = doc.styles["Normal"]
    if h.runs:
        h.runs[0].bold = True
        h.runs[0].font.size = docx_pt(18)
    for idx, art in enumerate(unmatched, 1):
        doc.add_page_break()
        title = doc.add_paragraph(f"Web-{idx}. {art.title}")
        title.style = doc.styles["Normal"]
        if title.runs:
            title.runs[0].bold = True
        for p in art.paras:
            par = doc.add_paragraph(p.text)
            par.style = doc.styles["Normal"]
            if p.kind == "lead" and par.runs:
                par.runs[0].bold = True


def docx_pt(value: int):
    from docx.shared import Pt

    return Pt(value)


def main():
    if not DOCX_PATH.is_file():
        raise FileNotFoundError(DOCX_PATH)
    missing = [name for name in HTML_FILES if not (HTML_ROOT / name).is_file()]
    if missing:
        raise FileNotFoundError("HTML lipsa: " + ", ".join(missing))

    doc = Document(DOCX_PATH)
    doc_articles = find_doc_articles(doc)
    html_articles = [parse_html_article(HTML_ROOT / name) for name in HTML_FILES]

    matches = []
    used_numbers: set[int] = set()
    unmatched: list[HtmlArticle] = []
    for html_article in html_articles:
        doc_article, score, mode = best_doc_match(html_article, doc_articles, used_numbers)
        if doc_article is None:
            unmatched.append(html_article)
            matches.append((html_article, None, score, mode))
        else:
            used_numbers.add(doc_article.number)
            matches.append((html_article, doc_article, score, mode))

    # Replace from the end so paragraph indices from the original analysis stay valid.
    for html_article, doc_article, score, mode in sorted(
        [m for m in matches if m[1] is not None], key=lambda item: item[1].marker_idx, reverse=True
    ):
        replace_article(doc, doc_article, html_article)

    append_unmatched(doc, unmatched)
    doc.save(OUT_DOCX)

    lines = []
    lines.append("Raport inlocuire articole web -> DOCX")
    lines.append(f"Document original: {DOCX_PATH}")
    lines.append(f"Document rezultat:  {OUT_DOCX}")
    lines.append(f"HTML-uri procesate: {len(html_articles)}")
    lines.append(f"Articole in carte:  {len(doc_articles)}")
    lines.append("")
    lines.append("INLOCUIRI")
    for html_article, doc_article, score, mode in matches:
        if doc_article is None:
            continue
        marker = "exact" if mode == "exact" else f"fuzzy {score:.2f}"
        lines.append(
            f"- {doc_article.number:02d}. {doc_article.title}  <=  {html_article.title}  [{marker}]  ({html_article.file.name})"
        )
    lines.append("")
    lines.append("HTML FARA CORESPONDENT SIGUR")
    if unmatched:
        for art in unmatched:
            lines.append(f"- {art.title} ({art.file.name})")
    else:
        lines.append("- niciunul")
    lines.append("")
    lines.append("NOTE")
    lines.append("- Originalul nu a fost modificat.")
    lines.append("- Cuprinsul static din document a fost pastrat; dupa finalizare, page numbers pot necesita actualizare in Word.")
    lines.append("- Articolele HTML nepotrivite sigur sunt anexate la final, daca exista.")
    REPORT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(OUT_DOCX)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
