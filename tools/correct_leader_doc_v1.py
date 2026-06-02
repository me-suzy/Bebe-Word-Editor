from __future__ import annotations

import re
from pathlib import Path
from typing import Callable

from docx import Document


BASE_DIR = Path(r"E:\Carte\BB\++++carti scrise de bebe\CELE 63 de calitati ale liderului")
INPUT_DOCX = BASE_DIR / "pentru tiparire - actualizat cu articole web.docx"
OUTPUT_DOCX = BASE_DIR / "Final Corectat V1.docx"
REPORT_TXT = BASE_DIR / "raport-corectura-v1.txt"


OLD_DIACRITICS = str.maketrans(
    {
        "ş": "ș",
        "Ş": "Ș",
        "ţ": "ț",
        "Ţ": "Ț",
    }
)


Edit = tuple[int, int, str, str]


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def apply_edits(paragraph, edits: list[Edit]) -> int:
    if not edits:
        return 0

    # Keep only non-overlapping edits, applied from right to left.
    edits = sorted(edits, key=lambda item: (item[0], item[1]))
    merged: list[Edit] = []
    last_end = -1
    for edit in edits:
        start, end, replacement, label = edit
        if start < last_end:
            continue
        merged.append(edit)
        last_end = end

    for start, end, replacement, _label in reversed(merged):
        runs = list(paragraph.runs)
        positions: list[tuple[int, int]] = []
        pos = 0
        for run in runs:
            run_len = len(run.text)
            positions.append((pos, pos + run_len))
            pos += run_len

        start_run = end_run = None
        for idx, (a, b) in enumerate(positions):
            if start_run is None and a <= start <= b:
                start_run = idx
            if end_run is None and a <= end <= b:
                end_run = idx
            if start_run is not None and end_run is not None:
                break

        if start_run is None or end_run is None:
            continue

        if start_run == end_run:
            a, _b = positions[start_run]
            run = runs[start_run]
            local_start = start - a
            local_end = end - a
            run.text = run.text[:local_start] + replacement + run.text[local_end:]
            continue

        first_a, _first_b = positions[start_run]
        last_a, _last_b = positions[end_run]
        first = runs[start_run]
        last = runs[end_run]
        first.text = first.text[: start - first_a] + replacement
        for idx in range(start_run + 1, end_run):
            runs[idx].text = ""
        last.text = last.text[end - last_a :]

    return len(merged)


def regex_edits(text: str, pattern: str, replacement: str | Callable[[re.Match], str], label: str, flags: int = 0) -> list[Edit]:
    edits: list[Edit] = []
    for match in re.finditer(pattern, text, flags):
        repl = replacement(match) if callable(replacement) else match.expand(replacement)
        if repl != match.group(0):
            edits.append((match.start(), match.end(), repl, label))
    return edits


def quote_edits(text: str) -> list[Edit]:
    edits: list[Edit] = []
    for idx, char in enumerate(text):
        if char != '"':
            continue
        prev = text[idx - 1] if idx > 0 else ""
        next_char = text[idx + 1] if idx + 1 < len(text) else ""
        if not prev or prev.isspace() or prev in "([{-:;":
            repl = "„"
        elif next_char and not next_char.isspace() and next_char not in ".,;:!?)]}":
            repl = "„"
        else:
            repl = "”"
        edits.append((idx, idx + 1, repl, "ghilimele drepte"))
    return edits


def is_marker(text: str) -> bool:
    stripped = text.strip()
    return len(stripped) >= 3 and stripped.startswith("-") and stripped.endswith("-") and stripped[1:-1].isdigit()


def main() -> None:
    doc = Document(INPUT_DOCX)
    stats: dict[str, int] = {}

    def bump(label: str, amount: int = 1) -> None:
        stats[label] = stats.get(label, 0) + amount

    # Same-length, run-local corrections first.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            original = run.text
            updated = original.translate(OLD_DIACRITICS)
            updated = updated.replace("“", "„")
            updated = updated.replace("\u00a0", " ")
            if updated != original:
                if original.translate(OLD_DIACRITICS) != original:
                    bump("diacritice vechi ş/ţ -> ș/ț")
                if "“" in original:
                    bump("ghilimele deschidere “ -> „")
                if "\u00a0" in original:
                    bump("spatii non-breaking normalizate")
                run.text = updated

    replacements: list[tuple[str, str, str, int]] = [
        (r"(?<![\wăâîșțĂÂÎȘȚ])si(?![\wăâîșțĂÂÎȘȚ])", "și", "si -> și", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])Si(?![\wăâîșțĂÂÎȘȚ])", "Și", "Si -> Și", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])in(?![\wăâîșțĂÂÎȘȚ])", "în", "in -> în", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])In(?![\wăâîșțĂÂÎȘȚ])", "În", "In -> În", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])Daca(?![\wăâîșțĂÂÎȘȚ])", "Dacă", "Daca -> Dacă", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])Dupa(?![\wăâîșțĂÂÎȘȚ])", "După", "Dupa -> După", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])Asa(?![\wăâîșțĂÂÎȘȚ])", "Așa", "Asa -> Așa", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])toata(?![\wăâîșțĂÂÎȘȚ])", "toată", "toata -> toată", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])raspuns(?![\wăâîșțĂÂÎȘȚ])", "răspuns", "raspuns -> răspuns", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])instructiuni(?![\wăâîșțĂÂÎȘȚ])", "instrucțiuni", "instructiuni -> instrucțiuni", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])mearga(?![\wăâîșțĂÂÎȘȚ])", "meargă", "mearga -> meargă", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])o singura(?![\wăâîșțĂÂÎȘȚ])", "o singură", "o singura -> o singură", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])decat(?![\wăâîșțĂÂÎȘȚ])", "decât", "decat -> decât", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])cate(?![\wăâîșțĂÂÎȘȚ])", "câte", "cate -> câte", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])cat(?![\wăâîșțĂÂÎȘȚ])", "cât", "cat -> cât", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])tau(?![\wăâîșțĂÂÎȘȚ])", "tău", "tau -> tău", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])ma bucur(?![\wăâîșțĂÂÎȘȚ])", "mă bucur", "ma bucur -> mă bucur", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])transferata(?![\wăâîșțĂÂÎȘȚ])", "transferată", "transferata -> transferată", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])supravietuit(?![\wăâîșțĂÂÎȘȚ])", "supraviețuit", "supravietuit -> supraviețuit", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])deajuns(?![\wăâîșțĂÂÎȘȚ])", "de ajuns", "deajuns -> de ajuns", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])niciodata(?![\wăâîșțĂÂÎȘȚ])", "niciodată", "niciodata -> niciodată", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])de-acord(?![\wăâîșțĂÂÎȘȚ])", "de acord", "de-acord -> de acord", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])Pe copii dumneavoastră(?![\wăâîșțĂÂÎȘȚ])", "Pe copiii dumneavoastră", "Pe copii -> Pe copiii", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])va face rău(?![\wăâîșțĂÂÎȘȚ])", "vă face rău", "va face rău -> vă face rău", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])în ce lor(?![\wăâîșțĂÂÎȘȚ])", "în ce loc", "în ce lor -> în ce loc", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])în cele din urma(?![\wăâîșțĂÂÎȘȚ])", "în cele din urmă", "în cele din urma -> în cele din urmă", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])asupra ce(?![\wăâîșțĂÂÎȘȚ])", "asupra a ce", "asupra ce -> asupra a ce", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])sa(?=-)", "să", "sa- -> să-", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])O sa(?![\wăâîșțĂÂÎȘȚ])", "O să", "O sa -> O să", 0),
        (r"(?<![\wăâîșțĂÂÎȘȚ])o sa(?![\wăâîșțĂÂÎȘȚ])", "o să", "o sa -> o să", 0),
        (
            r"(?<![\wăâîșțĂÂÎȘȚ])sa\s+(?=(găsești|fii|transformi|influențezi|nu-l|ieși|le|surprindă|încep|continui|cazi)\b)",
            "să ",
            "sa + verb -> să + verb",
            0,
        ),
        (r"aduce aportul cuvenit", "aduce contribuția cuvenită", "aduce aportul -> aduce contribuția", 0),
        (r"unei promisiuni încheiată", "unei promisiuni încheiate", "acord adjectival promisiune", 0),
        (r"pentru ca să", "ca să", "pentru ca să -> ca să", re.IGNORECASE),
        (r"premiza", "premisa", "premiza -> premisa", re.IGNORECASE),
        (r"destructiv", "distructiv", "destructiv -> distructiv", re.IGNORECASE),
        (r"perioadă de timp", "perioadă", "perioadă de timp -> perioadă", re.IGNORECASE),
        (r"și și", "și", "cuvant repetat: și și", 0),
        (r"ci dimpotrivă,", "ci, dimpotrivă,", "ci dimpotrivă -> ci, dimpotrivă", re.IGNORECASE),
        (r"cât mai perfectă", "cât mai desăvârșită", "cât mai perfectă -> cât mai desăvârșită", 0),
        (r"o imagine cât mai desăvârșită", "o imagine cât mai apropiată de desăvârșire", "imagine cât mai desăvârșită -> apropiată de desăvârșire", 0),
        (r"unul din cel mai străluciți lideri", "unul dintre cei mai străluciți lideri", "unul din cel mai străluciți -> unul dintre cei mai străluciți", 0),
        (r"acea forța puternică", "acea forță puternică", "acea forța -> acea forță", 0),
        (r"acea\s+”\s*imagine”", "acea „imagine”", "acea ” imagine” -> acea „imagine”", 0),
        (r"destul de vie pentru ca", "suficient de vie pentru ca", "destul de vie -> suficient de vie", 0),
        (r"stabilirea unui acord comun cu", "stabilirea unui acord cu", "acord comun -> acord", 0),
    ]

    punctuation_patterns: list[tuple[str, str, str, int]] = [
        (r"(?<!\.)\.\.(?!\.)", ".", "doua puncte consecutive", 0),
        (r",\s*etc\.\.", " etc.", "virgula inainte de etc + punct dublu", re.IGNORECASE),
        (r",\s*etc", " etc", "virgula inainte de etc", re.IGNORECASE),
        (r"\s+([,.;:!?])", r"\1", "spatiu inainte de punctuatie", 0),
        (r"([„])\s+", r"\1", "spatiu dupa ghilimea deschidere", 0),
        (r"\s+([”])", r"\1", "spatiu inainte de ghilimea inchidere", 0),
        (r"(^|[\s(:;,\-])”\s*(?=\w)", r"\1„", "ghilimea inchidere folosita ca deschidere", 0),
        (r"(?<=\S)[ ]{2,}(?=\S)", " ", "spatii multiple interne", 0),
    ]

    for paragraph in doc.paragraphs:
        if not paragraph.runs:
            continue

        text = paragraph_text(paragraph)
        if not text or is_marker(text):
            continue

        edits = quote_edits(text)
        count = apply_edits(paragraph, edits)
        if count:
            bump("ghilimele drepte -> românești", count)

        for pattern, repl, label, flags in replacements:
            text = paragraph_text(paragraph)
            edits = regex_edits(text, pattern, repl, label, flags)
            count = apply_edits(paragraph, edits)
            if count:
                bump(label, count)

        for pattern, repl, label, flags in punctuation_patterns:
            text = paragraph_text(paragraph)
            edits = regex_edits(text, pattern, repl, label, flags)
            count = apply_edits(paragraph, edits)
            if count:
                bump(label, count)

    doc.save(OUTPUT_DOCX)

    lines = [
        "Raport corectura V1",
        f"Document initial: {INPUT_DOCX}",
        f"Document rezultat: {OUTPUT_DOCX}",
        "",
        "Corecturi aplicate:",
    ]
    for label, count in sorted(stats.items(), key=lambda item: (-item[1], item[0])):
        lines.append(f"- {label}: {count}")
    REPORT_TXT.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(OUTPUT_DOCX)
    print(REPORT_TXT)
    for label, count in sorted(stats.items(), key=lambda item: (-item[1], item[0])):
        print(f"{count:5d}  {label}")


if __name__ == "__main__":
    main()
